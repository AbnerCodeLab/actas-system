import tkinter as tk
from tkinter import ttk, messagebox
from openpyxl import load_workbook
from docx import Document
from datetime import datetime
import os
from copy import deepcopy

RUTA_EXCEL = "data/organizaciones.xlsx"
RUTA_PLANTILLA = "templates/acta_template.docx"
CARPETA_SALIDA = "output"


def cargar_datos():
    try:
        wb = load_workbook(RUTA_EXCEL, data_only=True)

        def leer_hoja(nombre_hoja):
            ws = wb[nombre_hoja]
            encabezados = [cell.value for cell in ws[1]]
            datos = []

            for row in ws.iter_rows(min_row=2, values_only=True):
                registro = dict(zip(encabezados, row))
                datos.append(registro)

            return datos

        organizaciones = leer_hoja("organizaciones")
        representantes = leer_hoja("representantes")
        resoluciones = leer_hoja("resoluciones")

        organizaciones = [
            org for org in organizaciones
            if str(org.get("estado", "")).strip().lower() == "activo"
        ]

        representantes = [
            rep for rep in representantes
            if str(rep.get("estado", "")).strip().lower() == "vigente"
        ]

        resoluciones = [
            res for res in resoluciones
            if str(res.get("estado", "")).strip().lower() == "vigente"
        ]

        return organizaciones, representantes, resoluciones

    except Exception as e:
        messagebox.showerror("Error", f"No se pudo leer el Excel:\n{e}")
        return [], [], []


def limpiar(valor):
    if valor is None:
        return ""
    valor = str(valor).strip()

    if valor.endswith(".0"):
        valor = valor[:-2]

    return valor


def obtener_representante(codigo_centro, cargo):
    codigo_centro = limpiar(codigo_centro)
    cargo = limpiar(cargo).lower()

    for rep in representantes:
        codigo_rep = limpiar(rep.get("codigo_centro"))
        cargo_rep = limpiar(rep.get("cargo")).lower()
        estado_rep = limpiar(rep.get("estado")).lower()

        if codigo_rep == codigo_centro and cargo_rep == cargo and estado_rep == "vigente":
            nombres = limpiar(rep.get("nombres"))
            apellido_paterno = limpiar(rep.get("apellido_paterno"))
            apellido_materno = limpiar(rep.get("apellido_materno"))
            dni = limpiar(rep.get("dni"))

            nombre_completo = f"{nombres} {apellido_paterno} {apellido_materno}".strip()
            return nombre_completo, dni

    return "", ""


def obtener_resolucion(codigo_centro):
    for res in resoluciones:
        if str(res.get("codigo_centro")).strip() == str(codigo_centro).strip():
            return res
    return {}

def reemplazar_texto_en_celda(celda, buscar, reemplazo):
    for parrafo in celda.paragraphs:
        texto_completo = "".join(run.text for run in parrafo.runs)

        if buscar in texto_completo:
            texto_completo = texto_completo.replace(buscar, str(reemplazo))

            if parrafo.runs:
                parrafo.runs[0].text = texto_completo
                for run in parrafo.runs[1:]:
                    run.text = ""
            else:
                parrafo.add_run(texto_completo)

def reemplazar_texto(documento, datos):

    def reemplazar_en_parrafo(parrafo):
        texto_completo = "".join(run.text for run in parrafo.runs)

        if not any(clave in texto_completo for clave in datos):
            return

        for clave, valor in datos.items():
            texto_completo = texto_completo.replace(clave, str(valor))

        # Conserva el formato del primer run
        if parrafo.runs:
            parrafo.runs[0].text = texto_completo
            for run in parrafo.runs[1:]:
                run.text = ""

    # Párrafos normales
    for parrafo in documento.paragraphs:
        reemplazar_en_parrafo(parrafo)

    # Tablas
    for tabla in documento.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for parrafo in celda.paragraphs:
                    reemplazar_en_parrafo(parrafo)


def limpiar_nombre_archivo(texto):
    caracteres_invalidos = ['\\', '/', ':', '*', '?', '"', '<', '>', '|']
    for caracter in caracteres_invalidos:
        texto = texto.replace(caracter, "")
    return texto.strip()


def mostrar_datos(event=None):
    seleccion = combo_centros.get()

    for org in organizaciones:
        texto = f"{org['codigo_centro']} - {org['nombre_centro']}"

        if texto == seleccion:
            codigo = org["codigo_centro"]

            presidenta, dni_presidenta = obtener_representante(codigo, "Presidenta")
            secretaria, dni_secretaria = obtener_representante(codigo, "Secretaria")
            tesorera, dni_tesorera = obtener_representante(codigo, "Tesorera")
            resolucion = obtener_resolucion(codigo)

            info = f"""
CENTRO DE ATENCIÓN
Código: {org.get('codigo_centro', '')}
Nombre: {org.get('nombre_centro', '')}
Tipo: {org.get('tipo_centro', '')}
Modalidad: {org.get('modalidad', '')}
Dirección: {org.get('direccion', '')}
Distrito: {org.get('distrito', '')}
Provincia: {org.get('provincia', '')}
Departamento: {org.get('departamento', '')}
Usuarios: {org.get('numero_usuarios', '')}
RUOS: {org.get('ruos', '')}

JUNTA DIRECTIVA
Presidenta: {presidenta} - DNI {dni_presidenta}
Secretaria: {secretaria} - DNI {dni_secretaria}
Tesorera: {tesorera} - DNI {dni_tesorera}

RESOLUCIÓN
N° Resolución: {resolucion.get('numero_resolucion', '')}
Fecha emisión: {resolucion.get('fecha_emision', '')}
Inicio mandato: {resolucion.get('fecha_inicio_junta', '')}
Fin mandato: {resolucion.get('fecha_termino_mandato', '')}
"""
            label_info.config(text=info)
            return

#-----elegir plantilla-----
def obtener_plantilla_por_tipo(tipo_centro):
    tipo = str(tipo_centro).strip().lower()

    if "comedor popular" in tipo:
        return "templates/acta_comedor.docx"

    if "olla comun" in tipo:
        return "templates/acta_olla_comun.docx"

    return "templates/acta_template.docx"

#---generar tabla dinamica para junta directiva---
def obtener_cargos_por_tipo(tipo_centro):
    tipo = str(tipo_centro).strip().lower()

    if "comedor" in tipo:
        return [
            "Presidenta",
            "Vicepresidenta",
            "Secretaria",
            "Tesorera",
            "Almacenera"
        ]

    return [
        "Presidenta",
        "Vicepresidenta",
        "Secretaria",
        "Tesorera",
        "Almacenera",
        "Vocal",
        "Fiscal"
    ]

def escribir_en_celda(celda, texto):
    parrafo = celda.paragraphs[0]
    
    if parrafo.runs:
        parrafo.runs[0].text = str(texto)
    else:
        parrafo.add_run(str(texto))

def llenar_tabla_junta(documento, codigo_centro, tipo_centro):
    cargos = obtener_cargos_por_tipo(tipo_centro)

    for tabla in documento.tables:
        if len(tabla.columns) == 4:

            # La fila 0 es encabezado
            # La fila 1 debe ser la fila modelo con formato
            if len(tabla.rows) < 2:
                messagebox.showerror(
                    "Error",
                    "La tabla de Junta Directiva debe tener una fila modelo debajo del encabezado."
                )
                return

            fila_modelo = tabla.rows[1]._tr

            # Eliminar filas anteriores excepto encabezado y modelo
            while len(tabla.rows) > 2:
                tabla._element.remove(tabla.rows[2]._element)

            # Eliminar la fila modelo original luego de copiarla
            tabla._element.remove(tabla.rows[1]._element)

            for cargo in cargos:
                nueva_fila = deepcopy(fila_modelo)
                tabla._tbl.append(nueva_fila)

                fila = tabla.rows[-1].cells
                nombre, dni = obtener_representante(codigo_centro, cargo)

                reemplazar_texto_en_celda(fila[0], "{{cargo}}", cargo)
                reemplazar_texto_en_celda(fila[1], "{{nombre}}", nombre)
                reemplazar_texto_en_celda(fila[2], "{{dni}}", dni)
                reemplazar_texto_en_celda(fila[3], "{{asistencia}}", "Presente")

            break

#----generar acta word---


def generar_acta():
    seleccion = combo_centros.get()

    if not seleccion:
        messagebox.showwarning("Advertencia", "Seleccione un centro de atención.")
        return

    if not os.path.exists(RUTA_PLANTILLA):
        messagebox.showerror("Error", f"No se encontró la plantilla:\n{RUTA_PLANTILLA}")
        return

    os.makedirs(CARPETA_SALIDA, exist_ok=True)

    for org in organizaciones:
        texto = f"{org['codigo_centro']} - {org['nombre_centro']}"

        if texto == seleccion:
            codigo = org["codigo_centro"]

            presidenta, dni_presidenta = obtener_representante(codigo, "Presidenta")
            vicepresidenta, dni_vicepresidenta = obtener_representante(codigo, "Vicepresidenta")
            secretaria, dni_secretaria = obtener_representante(codigo, "Secretaria")
            tesorera, dni_tesorera = obtener_representante(codigo, "Tesorera")
            almacenera, dni_almacenera = obtener_representante(codigo, "Almacenera")
            vocal, dni_vocal = obtener_representante(codigo, "Vocal")
            fiscal, dni_fiscal = obtener_representante(codigo, "Fiscal")

            resolucion = obtener_resolucion(codigo)

            hoy = datetime.now()

            datos = {
                "{{codigo_centro}}": org.get("codigo_centro", ""),
                "{{nombre_centro}}": org.get("nombre_centro", ""),
                "{{tipo_centro}}": org.get("tipo_centro", ""),
                "{{modalidad}}": org.get("modalidad", ""),
                "{{direccion}}": org.get("direccion", ""),
                "{{distrito}}": org.get("distrito", ""),
                "{{provincia}}": org.get("provincia", ""),
                "{{departamento}}": org.get("departamento", ""),
                "{{numero_usuarios}}": org.get("numero_usuarios", ""),
                "{{ruos}}": org.get("ruos", ""),

                "{{presidenta}}": presidenta,
                "{{dni_presidenta}}": dni_presidenta,
                "{{vicepresidenta}}": vicepresidenta,
                "{{dni_vicepresidenta}}": dni_vicepresidenta,
                "{{secretaria}}": secretaria,
                "{{dni_secretaria}}": dni_secretaria,
                "{{tesorera}}": tesorera,
                "{{dni_tesorera}}": dni_tesorera,
                "{{almacenera}}": almacenera,
                "{{dni_almacenera}}": dni_almacenera,
                "{{vocal}}": vocal,
                "{{dni_vocal}}": dni_vocal,
                "{{fiscal}}": fiscal,
                "{{dni_fiscal}}": dni_fiscal,

                "{{numero_resolucion}}": resolucion.get("numero_resolucion", ""),
                "{{fecha_emision}}": resolucion.get("fecha_emision", ""),
                "{{fecha_inicio_junta}}": resolucion.get("fecha_inicio_junta", ""),
                "{{fecha_termino_mandato}}": resolucion.get("fecha_termino_mandato", ""),
                "{{archivo_resolucion}}": resolucion.get("archivo_resolucion", ""),

                "{{hora}}": hoy.strftime("%H:%M"),
                "{{dia}}": hoy.strftime("%d"),
                "{{mes}}": hoy.strftime("%m"),
                "{{anio}}": hoy.strftime("%Y"),
                "{{fecha_actual}}": hoy.strftime("%d/%m/%Y"),
                "{{hora_cierre}}": "____",
            }

            ruta_plantilla = obtener_plantilla_por_tipo(org.get("tipo_centro", ""))

            if not os.path.exists(ruta_plantilla):
                messagebox.showerror("Error", f"No se encontró la plantilla:\n{ruta_plantilla}")
                return

            documento = Document(ruta_plantilla)

            llenar_tabla_junta(
                documento,
                codigo,
                org.get("tipo_centro", "")
            )

            reemplazar_texto(documento, datos)

            nombre_limpio = limpiar_nombre_archivo(org.get("nombre_centro", "acta"))
            archivo_salida = os.path.join(CARPETA_SALIDA, f"ACTA_{codigo}_{nombre_limpio}.docx")

            documento.save(archivo_salida)

            messagebox.showinfo("Éxito", f"Acta generada correctamente:\n{archivo_salida}")
            return


# ================= INTERFAZ =================

root = tk.Tk()
root.title("Sistema de Generación de Actas - PCA")
root.geometry("750x600")

organizaciones, representantes, resoluciones = cargar_datos()

tk.Label(
    root,
    text="Sistema de Generación de Actas - PCA",
    font=("Arial", 15, "bold")
).pack(pady=10)

tk.Label(
    root,
    text="Seleccione Centro de Atención",
    font=("Arial", 11)
).pack(pady=5)

opciones = [
    f"{org['codigo_centro']} - {org['nombre_centro']}"
    for org in organizaciones
]

combo_centros = ttk.Combobox(root, values=opciones, width=90, state="readonly")
combo_centros.pack(pady=5)
combo_centros.bind("<<ComboboxSelected>>", mostrar_datos)

tk.Button(
    root,
    text="Generar Acta Word",
    command=generar_acta,
    bg="green",
    fg="white",
    font=("Arial", 11, "bold"),
    width=25
).pack(pady=15)

label_info = tk.Label(
    root,
    text="",
    justify="left",
    font=("Arial", 10),
    anchor="w"
)
label_info.pack(pady=10)

root.mainloop()

